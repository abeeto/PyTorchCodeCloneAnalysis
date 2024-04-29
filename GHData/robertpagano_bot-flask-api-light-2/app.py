from flask import Flask, request, redirect, url_for, flash, jsonify, send_file, render_template, request
# from flask_bootstrap import Bootstrap
from werkzeug.exceptions import Forbidden, HTTPException, NotFound, RequestTimeout, Unauthorized
import numpy as np
import pandas as pd
import pickle as p
import json
import io
import os
import glob

from pathlib import Path

from summarization.textsummarization import bert_sum, bert_sum_dynamic
from docx import Document
from docx.shared import Inches

import word_doc_gen

from linkcheck import flag_private_urls, flag_private_urls_to_dict

app = Flask(__name__)

## Testing out HTML
@app.route('/')
def form():
  return render_template('index.html')

## Handling Errors
@app.errorhandler(NotFound)
def page_not_found_handler(e: HTTPException):
    return render_template('404.html'), 404


@app.errorhandler(Unauthorized)
def unauthorized_handler(e: HTTPException):
    return render_template('401.html'), 401


@app.errorhandler(Forbidden)
def forbidden_handler(e: HTTPException):
    return render_template('403.html'), 403


@app.errorhandler(RequestTimeout)
def request_timeout_handler(e: HTTPException):
    return render_template('408.html'), 408
    
## summarizer using form data as an input, returns text summary
@app.route('/api/v1/resources/text/', methods=['POST'])
def summarize_from_text():
    
    data = request.form["data"]
    summary = bert_sum(data)

    return summary

## summarizer that pulls text from word doc, returns text summary with a dynamic length
@app.route('/api/v1/resources/document/summary/dynamic', methods=['GET', 'POST'])
def summarize_from_file_dynamic_length():
    
    f = request.files['data']
    f.save('datafile.docx')
    document = Document('datafile.docx')
    text =''
    for para in document.paragraphs[1:]:
        text+=para.text
    
    summary = bert_sum_dynamic(text)

    return summary

## this takes in a word file, summarizes text, adds summary to end, and returns document. Used downstream in transform route below
def transform():
    data = 'datafile.docx'
    document = Document(data)
    new = Document()
    text =''
    for para in document.paragraphs[1:]:
        text+=para.text
    new.add_paragraph(bert_sum(text))
    return new

## shows html interface, for now users can submit a document, and it will add a summary to the end
@app.route('/transform', methods=["GET","POST"])
def transform_view():
    f = request.files['data_file']
    section = request.form['section']
    f.save('datafile.docx')
    if not f:
        return "Please upload a word document"
    result = transform()
    result.save('result.docx')

    ## This uploads the file to the correct folder in sharepoint
    upload_sp(section)

    return send_file('result.docx', attachment_filename='new_file.docx')


## this takes in a word file, scrapes for links, checks links, and returns a json object of the table of results of link checker
@app.route('/api/v1/resources/document/links/json', methods=['POST'])
def check_links_to_json():
    
    data = request.files["link"]
    data.save('linkfile.docx')
    results = flag_private_urls_to_dict('linkfile.docx')
    
    return jsonify(results)

def remove_articles():
    files = glob.glob('articles/**/*')
    for f in files:
        os.remove(f)

@app.route('/api/v1/resources/document/docbuilder', methods=["POST"])
def build_docs():
    '''
    this takes in two arrays - one for files and one for filepaths. 
    It then creates a dictionary with the following structure:
    
    {article_name: {
        doc: (docx object)
        month: (month)
        section: (section)}
    }q
    this will then be used to create the master document files
    '''
    f = request.form['trig']
    docx_filepaths = glob.glob('articles/**/*.docx')
    txt_filepaths = glob.glob('articles/**/*.txt')
    print(docx_filepaths, txt_filepaths)
    word_doc_gen.make_master_file(docx_filepaths, txt_filepaths)
    # remove_articles()
    return send_file('master_doc.docx', attachment_filename='new_file.docx')

@app.route('/api/v1/resources/document/save', methods=['POST'])
def save_articles():
    data = request.files["file"]
    section = request.form["section"]
    filepath = data.filename
    title = filepath.split('/')[-1]
    path = f'articles/{section}/{title}'
    Path(f'articles/{section}').mkdir(parents=True, exist_ok=True)
    data.save(path)
    return ""

if __name__ == '__main__':
    app.run(debug=True)
