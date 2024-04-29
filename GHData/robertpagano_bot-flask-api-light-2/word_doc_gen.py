from docx import Document

import pandas as pd

import glob

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from docxcompose.composer import Composer

from pathlib import Path

'''
functions extract all needed data from filepaths_to_docs list and filepaths_to_summs list
use this data to create a dataframe, and then build the master document 
'''

def make_doc_dataframe(filepaths_to_docs, filepaths_to_summs):

    '''
    this takes in two lists:
    filepaths_to_docs - filepaths to all of the articles in word doc form
    filepaths_to_summs - approved summaries of the articles, in txt form with "name + '_summary'" naming convention
    this function parses the paths in filepaths_to_docs and text in filepaths_to_summs to create a dataframe that
    is used to build the master document. this dataframe is sorted by section, and can be ordered in any section order
    by changing the value of 'section_order'.
    the columns in the dataframe are:
    
        - article_name - parsed article_name from filepath
        - section - parsed section content from filepath
        - doc_object - python docx object read in from filepath
        - summ_text - text extracted from each summary file
        - new_section - True if it's the first row of a specific section
    '''

    cols = ['article_name', 'section', 'doc_object', 'summ_text'] # new_section column created below

    temp_list = []

    # each list will be ['article_name', 'section', 'doc_object', 'summ_text'] for one path pair, 
    # then be appended to dataframe

    df = pd.DataFrame(columns=cols)

    # zip_dict = dict(zip(filepaths_list, zip(doc_list, summ_list_txt)))
    zipped_paths_dict = dict(zip(filepaths_to_docs, filepaths_to_summs))

    for doc_path, txt_path  in zipped_paths_dict.items():

        article_name = doc_path.split('\\')[-1][:-5]
        section = doc_path.split('\\')[1]
        article_doc = Document(doc_path)
        summ_text = Path(txt_path).read_text()

        temp_list.append(article_name)
        temp_list.append(section)
        temp_list.append(article_doc)
        temp_list.append(summ_text)

        # convert list to series using cols, then append series to dataframe
        temp_series = pd.Series(temp_list, index = df.columns)
        df = df.append(temp_series, ignore_index = True)
        temp_list = [] # empty the list for the next iteration

    # 'section_order' is the order we want the sections to be in the final document
    # this sorts the dataframe by a specific order in sections    
    section_order = ["Content + Training", "Product + Availability", "Programs + Offers", "Partner Update", "nocat"]
    df['section'] = pd.Categorical(df['section'], section_order)

    df = df.sort_values('section')

    # create new_section column to indicate when we need a section heading in master doc
    df['new_section'] = df['section'].shift().fillna('nocat') != df['section']

    return df

'''
TO DO - MORNING OF AUGUST 20
  - make a document for each heading that lives in the repo, can pull from these using composer to add to the newsletter
- use font information that maria sent
  - All titles: Segoe UI Semibold font size 12 (black) 
    body/paragraphs - segoe ui font size 11 (black)   	
    unfilled bullets (like the photo - black) 
    Abstract/article : segoe ui semibold font size 11
- make sure in landscape mode (sideways)
'''

def make_toc(doc):
    '''
    this function creates a table of contents object within a docx object, which will be called when the master files are created
    it indexes any text with "heading styles"
    if article "sections" are heading 1, and article titles are heading 2, it will take care of all of the proper formatting/indenting
    '''
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
    return doc

def create_doc(doc, summ_text, section, article_name, new_section = False):
    '''
    if new_section == True, then this will make a heading. This will only get triggered as True when a new section is
    seen in the dataframe
    '''
    
    try:
        doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    except:
        print('style already added')
        pass

    paragraphs = doc.paragraphs

    if new_section == True:

        p = paragraphs[0]
        section_paragraph = p.insert_paragraph_before(section)
        section_paragraph.style = doc.styles['Heading 1']
        title = p.insert_paragraph_before(article_name)
        title.style = doc.styles['Heading 2']
        abstract_title = p.insert_paragraph_before('Abstract:')
        abstract = p.insert_paragraph_before(summ_text)
        article_title = p.insert_paragraph_before('Article:')

    else:

        p = paragraphs[0]
        title = p.insert_paragraph_before(article_name)
        title.style = doc.styles['Heading 2']
        abstract_title = p.insert_paragraph_before('Abstract:')
        abstract = p.insert_paragraph_before(summ_text)
        article_title = p.insert_paragraph_before('Article:')

    return doc

def make_master_file(filepaths_to_docs, filepaths_to_summs):

    '''
    takes in doc filepaths and summ filepaths (which will be the inputs from the API function)
    returns the master file by calling the already defined functions
    '''

    df = make_doc_dataframe(filepaths_to_docs, filepaths_to_summs)

    toc = Document()
    paragraph = toc.add_paragraph('TABLE OF CONTENTS')
    toc = make_toc(toc)
    article_list = [toc]

    for index, row in df.iterrows():
        doc = row['doc_object']
        summ_text = row['summ_text']
        section = row['section']
        article_name = row['article_name']
        
        if row['new_section'] == True: # this adds the section heading above first article in section
            article_list.append(create_doc(doc, summ_text, section, article_name, new_section = True))

        else:
            article_list.append(create_doc(doc, summ_text, section, article_name, new_section = False))

    master = article_list[0]
    composer = Composer(master)
    for document in article_list[1:]:
        composer.append(document)

    composer.save('master_doc.docx') # need to tell Kevin that I'm saving it down as master_doc.docx

    # return master_doc 

